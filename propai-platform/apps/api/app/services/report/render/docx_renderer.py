"""DOCX 렌더러 — 정본 ReportModel 을 Block 단위로 순회하며 python-docx 문서 생성.

표지 → (핵심요약) → 섹션들(블록 순회) → 면책. pdf_renderer 와 '같은 분기·같은 순서'로 미러링한다.
씨앗: market_report_service.to_docx 의 _h/_table/_caption 헬퍼를 'Block 구동 일반형'으로 승격.

★순수 모듈: reportlab/FastAPI/DB 임포트 금지. python-docx + 표준 라이브러리만.
★가짜값 금지: 값은 model.fmt_value 로만 문자열화(없으면 '—'). 빈 표는 '데이터 없음' 정직 표기.
★docx 는 셀 배경·차트가 네이티브로 약함 → 셀 배경은 w:shd XML 로 직접, 차트는 표로 폴백.
"""

from __future__ import annotations

import contextlib
import io
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor

from . import tokens as T
from .model import ReportModel, Section, fmt_value


# ── 저수준 헬퍼(토큰 → python-docx 변환) ─────────────────────────────
def _rgb(hex_str: str) -> RGBColor:
    """'#0e7490' → docx RGBColor. 토큰 HEX 를 docx 색 객체로 바꾼다."""
    r, g, b = T.hex_to_rgb(hex_str)
    return RGBColor(r, g, b)


def _safe_rgb(hex_str: str | None, default: str) -> RGBColor:
    """색 문자열이 이상하면(형식 오류) 기본색으로 안전 폴백."""
    try:
        if hex_str:
            return _rgb(hex_str)
    except (ValueError, IndexError, TypeError):
        pass
    return _rgb(default)


def _pt(key: str) -> Pt:
    """토큰 타이포 스케일 → docx Pt. 본문/표는 화면 뷰어 가독을 위해 +1(토큰 주석)."""
    size = T.TYPE[key][0]
    if key in ("body", "table", "table_header"):
        size += 1
    return Pt(size)


def _set_kr_font(run) -> None:
    """글자에 한글 폰트(맑은 고딕) 지정. docx 는 라틴·한글 폰트가 따로라 eastAsia 도 같이 박아준다."""
    run.font.name = T.FONT_KR_GOTHIC  # 라틴(ascii/hAnsi) 이름
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), T.FONT_KR_GOTHIC)  # 한글은 eastAsia 로 지정해야 적용됨


def _shade_cell(cell, hex_color: str) -> None:
    """표 셀 배경색 칠하기. docx 는 셀 배경이 네이티브 API 로 없어 w:shd XML 을 직접 넣는다."""
    hexv = (hex_color or "").lstrip("#")
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexv)
    tc_pr.append(shd)


def add_hyperlink(paragraph, url: str, text: str, color_hex: str = T.LINK):
    """문단에 진짜 클릭되는 하이퍼링크 넣기(법령 링크용). 파란 밑줄 + 한글 폰트."""
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)  # 외부 URL 관계 등록
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), (color_hex or T.LINK).lstrip("#"))
    rpr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:eastAsia"), T.FONT_KR_GOTHIC)
    rpr.append(rfonts)
    new_run.append(rpr)
    t_el = OxmlElement("w:t")
    t_el.text = text
    new_run.append(t_el)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def _apply_grid(table) -> None:
    """표에 얇은 격자선. 배경색은 우리가 w:shd 로 직접 칠하므로, 색을 안 덮는 'Table Grid' 사용."""
    with contextlib.suppress(Exception):  # 스타일 없으면 기본 유지
        table.style = "Table Grid"


def _run(para, text: str, *, bold: bool = False, size_key: str = "body",
         color_hex: str | None = None):
    """문단에 글자(run) 하나 추가 + 한글 폰트·크기·색 지정. 텍스트는 fmt_value 로만."""
    run = para.add_run(fmt_value(text))
    run.font.bold = bold
    run.font.size = _pt(size_key)
    _set_kr_font(run)
    run.font.color.rgb = _rgb(color_hex) if color_hex else _rgb(T.INK)
    return run


def _para(doc, text: str, *, bold: bool = False, size_key: str = "body",
          color_hex: str | None = None):
    """한 줄 문단 추가(공용). 캡션·메타·본문 모두 이걸로."""
    p = doc.add_paragraph()
    _run(p, text, bold=bold, size_key=size_key, color_hex=color_hex)
    return p


def _muted(doc, text: str, size_key: str = "caption"):
    """작고 흐린 문단(캡션·출처·면책·'데이터 없음')."""
    return _para(doc, text, size_key=size_key, color_hex=T.MUTED)


def _fill_cell(cell, text: str, *, bold: bool = False, color_hex: str | None = None,
               align=None, shade_hex: str | None = None) -> None:
    """표 셀에 값 채우기 + 정렬·굵기·글자색·배경 shade. 값은 fmt_value 로만(가짜 금지)."""
    para = cell.paragraphs[0]
    run = para.add_run(fmt_value(text))
    run.font.bold = bold
    run.font.size = _pt("table")
    _set_kr_font(run)
    if color_hex:
        run.font.color.rgb = _rgb(color_hex)
    if align is not None:
        para.alignment = align
    if shade_hex:
        _shade_cell(cell, shade_hex)


def _h(doc, text: str, level: int = 1):
    """섹션 제목(딥틸색). level=0 은 표지 대제목, level=1 은 섹션 제목."""
    heading = doc.add_heading(fmt_value(text), level=level)
    for run in heading.runs:
        run.font.color.rgb = _rgb(T.BRAND)
        _set_kr_font(run)
    return heading


def _h3(doc, text: str):
    """블록 소제목(표/차트/근거 제목). 잉크색 굵게."""
    return _para(doc, text, bold=True, size_key="h3", color_hex=T.INK)


# ── 표 블록 헬퍼 ─────────────────────────────────────────────────────
def _kv_table(doc, rows: list[tuple[str, Any]]):
    """키-값 2열표. 라벨셀은 zebra 배경·굵게, 값셀은 일반."""
    table = doc.add_table(rows=0, cols=2)
    _apply_grid(table)
    for label, value in rows:
        cells = table.add_row().cells
        _fill_cell(cells[0], label, bold=True, shade_hex=T.ZEBRA)
        _fill_cell(cells[1], value)
    doc.add_paragraph()  # 표 뒤 간격
    return table


def _data_table(doc, headers: list[str], rows: list[list[Any]],
                numeric_cols: list[int] | None = None, total_row: bool = False):
    """일반 데이터표. 헤더=딥틸 배경 흰 굵게, 짝수행 zebra shade, 숫자열 우측정렬, 합계행 굵게."""
    numeric = set(numeric_cols or [])
    ncols = len(headers)
    table = doc.add_table(rows=1, cols=ncols)
    _apply_grid(table)
    # 헤더행: 딥틸 배경 + 흰색 굵게
    for c, htext in enumerate(headers):
        _fill_cell(table.rows[0].cells[c], htext, bold=True,
                   color_hex=T.WHITE, shade_hex=T.BRAND)
    last = len(rows) - 1
    for ri, row in enumerate(rows):
        cells = table.add_row().cells
        is_total = total_row and ri == last              # 마지막 행을 합계로
        zebra = (ri % 2 == 1) and not is_total           # 첫 본문행=흰, 다음=zebra 교대
        for c in range(ncols):
            val = row[c] if c < len(row) else None
            align = WD_ALIGN_PARAGRAPH.RIGHT if c in numeric else None
            shade = T.PANEL if is_total else (T.ZEBRA if zebra else None)
            _fill_cell(cells[c], val, bold=is_total, align=align, shade_hex=shade)
    doc.add_paragraph()
    return table


def _kpi_row(doc, tiles: list):
    """KPI 타일 행(1행 N열 표). 각 칸=라벨(작게)+수치(크게·signal색)+기준(작게)."""
    n = max(1, len(tiles))
    table = doc.add_table(rows=1, cols=n)
    _apply_grid(table)
    for i, tile in enumerate(tiles):
        cell = table.rows[0].cells[i]
        # 라벨(가운데·흐림)
        p_label = cell.paragraphs[0]
        p_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rl = p_label.add_run(fmt_value(tile.label))
        _set_kr_font(rl)
        rl.font.size = _pt("kpi_label")
        rl.font.color.rgb = _rgb(T.MUTED)
        # 수치(가운데·크게·굵게·임계 신호색)
        p_val = cell.add_paragraph()
        p_val.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rv = p_val.add_run(fmt_value(tile.value))
        _set_kr_font(rv)
        rv.font.bold = True
        rv.font.size = _pt("kpi_value")
        rv.font.color.rgb = _safe_rgb(tile.signal, T.INK) if tile.signal else _rgb(T.INK)
        # 기준(선택·작게)
        if tile.basis:
            p_b = cell.add_paragraph()
            p_b.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rb = p_b.add_run(fmt_value(tile.basis))
            _set_kr_font(rb)
            rb.font.size = _pt("caption")
            rb.font.color.rgb = _rgb(T.MUTED)
    doc.add_paragraph()
    return table


def _grade_badge(doc, grade: str, label: str | None):
    """등급 배지 = 1×1 shade 셀(tint 배경 + 진한 등급색 글자)."""
    gs = T.grade_style(grade)
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.rows[0].cells[0]
    _shade_cell(cell, gs["bg"])
    try:  # 배지는 좁게(전체폭 금지)
        table.columns[0].width = Cm(6)
        cell.width = Cm(6)
    except Exception:  # noqa: BLE001
        pass
    para = cell.paragraphs[0]
    if label:
        r0 = para.add_run(fmt_value(label) + "  ")
        _set_kr_font(r0)
        r0.font.size = _pt("body")
        r0.font.color.rgb = _rgb(gs["fg"])
    r1 = para.add_run(fmt_value(gs["label"]))
    _set_kr_font(r1)
    r1.font.bold = True
    r1.font.size = _pt("body")
    r1.font.color.rgb = _rgb(gs["fg"])
    doc.add_paragraph()
    return table


# ── 블록 디스패치(pdf_renderer 와 동일 분기·순서) ─────────────────────
def _render_block(doc, block: Any) -> None:
    """Block.kind 별 렌더 디스패치."""
    kind = getattr(block, "kind", None)

    if kind == "kv":
        if block.title:
            _h3(doc, block.title)
        if block.rows:
            _kv_table(doc, block.rows)
        else:
            _muted(doc, T.EMPTY_MARK)

    elif kind == "table":
        if block.title:
            _h3(doc, block.title)
        if block.rows:
            _data_table(doc, block.headers, block.rows,
                        numeric_cols=block.numeric_cols, total_row=block.total_row)
        else:
            _muted(doc, "데이터 없음")
        if block.caption:
            _muted(doc, block.caption)

    elif kind == "kpi":
        if block.tiles:
            _kpi_row(doc, block.tiles)

    elif kind == "chart":
        # docx 는 네이티브 차트가 없어 categories×series 를 표로 폴백(가짜 곡선 금지).
        _h3(doc, block.title)
        if block.series:
            headers = ["구분", *[s.name for s in block.series]]
            rows = []
            for i in range(len(block.categories)):
                cat = block.categories[i] if i < len(block.categories) else T.EMPTY_MARK
                row: list[Any] = [cat]
                for s in block.series:
                    row.append(s.values[i] if i < len(s.values) else None)
                rows.append(row)
            _data_table(doc, headers, rows, numeric_cols=list(range(1, len(headers))))
        else:
            _muted(doc, "데이터 없음")
        if block.caption:
            _muted(doc, block.caption)

    elif kind == "narrative":
        if block.title:
            _h3(doc, block.title)
        for para in block.paragraphs or []:
            if str(para).strip():
                _para(doc, para)

    elif kind == "evidence":
        if block.title:
            _h3(doc, block.title)
        for ev in block.items or []:
            conf = str(ev.confidence).lower() if ev.confidence else ""
            p = doc.add_paragraph()
            # 값(굵게)
            rv = p.add_run(fmt_value(ev.value))
            _set_kr_font(rv)
            rv.font.bold = True
            rv.font.size = _pt("caption")
            rv.font.color.rgb = _rgb(T.INK)
            # 근거·출처
            tail = []
            if ev.basis:
                tail.append(f"근거: {fmt_value(ev.basis)}")
            if ev.source:
                tail.append(f"출처: {fmt_value(ev.source)}")
            if tail:
                rt = p.add_run(" · " + " · ".join(tail))
                _set_kr_font(rt)
                rt.font.size = _pt("caption")
                rt.font.color.rgb = _rgb(T.MUTED)
            # 저신뢰(low/med)는 앰버로 명확히 표기(R4)
            if conf in ("low", "med", "medium"):
                ra = p.add_run(f" · (신뢰도 {fmt_value(ev.confidence)})")
                _set_kr_font(ra)
                ra.font.size = _pt("caption")
                ra.font.color.rgb = _rgb(T.AMBER)
            # 법령 링크는 verified/high(또는 미표기)만 노출
            if ev.legal_link and (conf in ("high", "verified") or ev.confidence is None):
                rs = p.add_run(" · ")
                _set_kr_font(rs)
                rs.font.size = _pt("caption")
                rs.font.color.rgb = _rgb(T.MUTED)
                add_hyperlink(p, ev.legal_link, "법령", T.LINK)

    elif kind == "checklist":
        if block.title:
            _h3(doc, block.title)
        rows = []
        for label, status in block.items or []:
            mark = "✓" if status is True else ("—" if status in (False, None) else fmt_value(status))
            rows.append([label, mark])
        if rows:
            _data_table(doc, ["항목", "상태"], rows)

    elif kind == "grade":
        _grade_badge(doc, block.grade, block.label)

    elif kind == "image":
        try:
            width = Mm(block.max_width_mm or 150)
            doc.add_picture(io.BytesIO(block.png), width=width)
        except Exception:  # noqa: BLE001  # 이미지 깨져도 문서는 계속
            _muted(doc, "(이미지 없음)")
        if block.caption:
            _muted(doc, block.caption)

    elif kind == "disclaimer":
        _muted(doc, block.text, size_key="disclaimer")


def _render_section(doc, sec: Section) -> None:
    """섹션 제목 + 블록들(pdf_renderer._render_section 미러)."""
    heading = f"{sec.section_no}. {sec.title}" if sec.section_no else sec.title
    _h(doc, heading, level=1)
    for block in sec.blocks:
        _render_block(doc, block)


# ── 문서 기본 스타일(폰트·여백) ──────────────────────────────────────
def _apply_base_style(doc) -> None:
    """기본(Normal) 스타일에 한글 폰트·본문 크기 + 페이지 여백을 건다."""
    style = doc.styles["Normal"]
    style.font.name = T.FONT_KR_GOTHIC
    style.font.size = _pt("body")
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), T.FONT_KR_GOTHIC)
    margin = Cm(T.DOCX["margin_cm"])
    for section in doc.sections:
        section.top_margin = margin
        section.bottom_margin = margin
        section.left_margin = margin
        section.right_margin = margin


def _cover(doc, meta) -> None:
    """표지: 대제목(강조색) + 메타 문단(주소·생성일·문서번호·대외비·데이터 채움도)."""
    title = doc.add_heading(fmt_value(meta.title), level=0)
    accent = meta.accent_color or T.BRAND
    for run in title.runs:
        run.font.color.rgb = _safe_rgb(accent, T.BRAND)
        _set_kr_font(run)
    if meta.subtitle:
        _para(doc, meta.subtitle, size_key="body", color_hex=T.SECONDARY)
    if meta.project_address:
        _para(doc, f"대상지: {fmt_value(meta.project_address)}")
    if meta.generated_at:
        _para(doc, f"생성일: {fmt_value(meta.generated_at)}")
    if meta.doc_no:
        _para(doc, f"문서번호: {fmt_value(meta.doc_no)}")
    _para(doc, T.BRANDING, color_hex=T.MUTED)
    if meta.confidential:
        _para(doc, T.CONFIDENTIAL_LABEL, bold=True, color_hex=T.SIGNAL["danger"])
    if meta.completeness:
        pct = meta.completeness.get("pct")
        if pct is not None:
            _muted(doc, f"데이터 채움도: {fmt_value(pct)}% (정직 표기)")
    doc.add_paragraph()


# ── 최상위 진입점 ────────────────────────────────────────────────────
def render_docx(model: ReportModel) -> bytes:
    """정본 모델 → DOCX bytes. 표지 → 핵심요약 → 섹션들 → 면책 순(세 포맷 동일 구조)."""
    doc = Document()
    _apply_base_style(doc)

    # ── 표지 ──
    _cover(doc, model.meta)

    # ── 핵심 요약(두괄식) ──
    if model.exec_summary:
        _render_section(doc, model.exec_summary)

    # ── 섹션들 ──
    for sec in model.sections:
        _render_section(doc, sec)

    # ── 면책 ──
    doc.add_paragraph()
    _muted(doc, model.disclaimer or T.DISCLAIMER_TEXT, size_key="disclaimer")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
